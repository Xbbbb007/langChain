import os
from typing import List
from pypdf import PdfReader
from docx import Document
import openpyxl
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LC_Document

class FileParser:
    @staticmethod
    def parse_pdf(file_path: str) -> str:
        text = ""
        reader = PdfReader(file_path)
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text += f"\n[第 {i + 1} 页]\n" + page_text
        return text

    @staticmethod
    def parse_docx(file_path: str) -> str:
        doc = Document(file_path)
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                text.append(" | ".join(row_data))
        return "\n".join(text)

    @staticmethod
    def parse_excel(file_path: str) -> str:
        wb = openpyxl.load_workbook(file_path, data_only=True)
        text = []
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text.append(f"\n[工作表: {sheet_name}]\n")
            
            # Get headers from first row
            rows = list(sheet.iter_rows(values_only=True))
            if not rows:
                continue
            
            headers = rows[0]
            # Replace None in headers
            headers = [str(h) if h is not None else f"列{idx}" for idx, h in enumerate(headers)]
            
            for row_idx, row in enumerate(rows[1:], start=2):
                if not any(v is not None for v in row):
                    continue  # skip empty rows
                row_items = []
                for col_idx, val in enumerate(row):
                    if val is not None:
                        # Format header: value
                        h = headers[col_idx] if col_idx < len(headers) else f"列{col_idx}"
                        row_items.append(f"{h}: {val}")
                if row_items:
                    text.append(f"行 {row_idx}: " + ", ".join(row_items))
        return "\n".join(text)

    @staticmethod
    def parse_txt(file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

def load_and_split_document(file_path: str, filename: str, document_id: str) -> List[LC_Document]:
    ext = os.path.splitext(filename)[1].lower()
    
    if ext == ".pdf":
        raw_text = FileParser.parse_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        raw_text = FileParser.parse_docx(file_path)
    elif ext in [".xlsx", ".xls"]:
        raw_text = FileParser.parse_excel(file_path)
    elif ext in [".txt", ".md", ".json"]:
        raw_text = FileParser.parse_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")

    # Split using LangChain RecursiveCharacterTextSplitter
    # Standard chunk sizes: 500 characters with 100 overlap (good for Chinese)
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=600,
        chunk_overlap=120,
        length_function=len,
        separators=["\n\n", "\n", "。", "！", "？", "；", " ", ""]
    )
    
    chunks = text_splitter.split_text(raw_text)
    
    documents = []
    for idx, chunk in enumerate(chunks):
        doc = LC_Document(
            page_content=chunk,
            metadata={
                "document_id": document_id,
                "source_name": filename,
                "chunk_id": idx,
                "text": chunk  # store in metadata too for ease of display if needed
            }
        )
        documents.append(doc)
        
    return documents
